#!/usr/bin/env python3
"""DOCX 양식에 섹션 내용 자동 기입 (python-docx).

HWP 양식은 soffice로 DOCX 변환 후 이 스크립트로 기입, 다시 HWP로 재변환.

사용법:
    python fill_docx.py <template.docx> <content_dir> [--output <output.docx>]
"""
import sys
import json
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def load_section_mapping(content_dir: str) -> dict:
    """fill-guide/section-mapping.json 로드."""
    mapping_path = Path(content_dir) / "fill-guide" / "section-mapping.json"
    if mapping_path.exists():
        return json.loads(mapping_path.read_text(encoding="utf-8"))
    # 폴백: drafts/ 폴더에서 섹션 파일 자동 매핑
    drafts = Path(content_dir) / "drafts"
    sections = {}
    if drafts.exists():
        for f in sorted(drafts.glob("section-*.md")):
            match = re.match(r"section-(\d+)-(.+)\.md", f.name)
            if match:
                sections[f"section-{match.group(1)}"] = {
                    "file": str(f),
                    "title": match.group(2).replace("-", " "),
                }
    return {"sections": sections}


def find_and_replace_placeholder(doc, placeholder: str, content: str):
    """문서에서 플레이스홀더 텍스트를 찾아 내용으로 교체."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, content)
            return True
    # 테이블 셀도 확인
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, content)
                    return True
    return False


def fill_by_heading(doc, heading_text: str, content: str):
    """특정 헤딩 아래에 내용 삽입."""
    found = False
    for i, para in enumerate(doc.paragraphs):
        if found:
            # 다음 헤딩 전까지의 빈 단락에 내용 삽입
            if para.style.name.startswith("Heading"):
                break
            if not para.text.strip():
                para.text = content
                para.style = doc.styles["Normal"]
                return True
        if heading_text.lower() in para.text.lower():
            found = True
    return False


def fill_document(template_path: str, content_dir: str, output_path: str = None):
    """메인 기입 함수."""
    doc = Document(template_path)
    mapping = load_section_mapping(content_dir)
    filled_count = 0

    for section_id, section_info in mapping.get("sections", {}).items():
        content_file = Path(section_info.get("file", ""))
        if not content_file.exists():
            print(f"  SKIP: {section_id} — file not found: {content_file}")
            continue

        content = content_file.read_text(encoding="utf-8")
        # 마크다운 헤더 제거 (양식에는 이미 헤더가 있음)
        content = re.sub(r"^#{1,6}\s+.*$", "", content, flags=re.MULTILINE).strip()

        title = section_info.get("title", section_id)
        if fill_by_heading(doc, title, content):
            filled_count += 1
            print(f"  FILL: {section_id} → '{title}'")
        else:
            print(f"  MISS: {section_id} — heading '{title}' not found in template")

    if not output_path:
        output_path = template_path.replace(".docx", "-filled.docx")

    doc.save(output_path)
    print(f"\nSaved: {output_path} ({filled_count} sections filled)")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python fill_docx.py <template.docx> <content_dir> [--output <path>]")
        sys.exit(1)

    template = sys.argv[1]
    content = sys.argv[2]
    output = None
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output = sys.argv[idx + 1]

    fill_document(template, content, output)
