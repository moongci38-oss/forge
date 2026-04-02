#!/usr/bin/env python3
"""
md-to-docx.py — Markdown → DOCX 프로그래매틱 변환 스크립트
pandoc 블랙박스 변환 대신 python-docx로 요소별 직접 빌드.

사용법:
  python3 md-to-docx.py input.md output.docx
  python3 md-to-docx.py --dir /path/to/full-text/  # 폴더 내 모든 .md 변환

특징:
  - 멀티라인 색상 span (red/blue) → RGBColor 직접 매핑
  - 이미지 ![](path) → InlineShape 바이너리 임베딩
  - 표 → Table 객체 직접 생성
  - 제목 ##/###/#### → Heading 1/2/3
  - 한글 폰트: 맑은 고딕 (w:eastAsia 명시)
  - --- → 페이지 브레이크
  - 문단 간격 최소화 (space_after=Pt(2))
"""

import re
import sys
import os
from pathlib import Path
from PIL import Image as PILImage

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ─── 폰트 설정 ───

FONT_TITLE = "맑은 고딕"
FONT_BODY = "맑은 고딕"
FONT_TABLE = "맑은 고딕"
FONT_FALLBACK = "Malgun Gothic"

TITLE_SIZE = Pt(18)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(11)
BODY_SIZE = Pt(11)
TABLE_SIZE = Pt(10)
CAPTION_SIZE = Pt(9)


def set_font(run, font_name=FONT_BODY, size=BODY_SIZE, bold=False, color=None):
    """run에 폰트, 크기, 색상을 설정한다. w:eastAsia로 한글 폰트를 명시."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, space_before=Pt(0), space_after=Pt(2)):
    """문단 간격을 설정. 줄간격 120%, 뒤 2pt (간격 최소화)."""
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing = 1.2


def set_style_defaults(doc):
    """문서 기본 스타일을 맑은 고딕으로 설정."""
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.2
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_BODY)

    for level, (hsize, hbold) in enumerate(
        [(H1_SIZE, True), (H2_SIZE, True), (H3_SIZE, True)], start=1
    ):
        hstyle_name = f"Heading {level}"
        if hstyle_name in doc.styles:
            hs = doc.styles[hstyle_name]
            hs.font.name = FONT_TITLE
            hs.font.size = hsize
            hs.font.bold = hbold
            hs.paragraph_format.space_before = [Pt(12), Pt(10), Pt(8)][level-1]
            hs.paragraph_format.space_after = [Pt(6), Pt(4), Pt(4)][level-1]
            hrpr = hs.element.get_or_add_rPr()
            hrf = hrpr.find(qn("w:rFonts"))
            if hrf is None:
                hrf = hs.element.makeelement(qn("w:rFonts"), {})
                hrpr.insert(0, hrf)
            hrf.set(qn("w:eastAsia"), FONT_TITLE)


# ─── 색상 파싱 ───

COLOR_MAP = {
    "red": RGBColor(0xFF, 0x00, 0x00),
    "blue": RGBColor(0x00, 0x62, 0xB8),
    "#ff0000": RGBColor(0xFF, 0x00, 0x00),
    "#0062b8": RGBColor(0x00, 0x62, 0xB8),
}

# 인라인 span (한 줄 안에서 열고 닫는 경우)
INLINE_SPAN_RE = re.compile(
    r'<span\s+style="color:\s*([^";]+)[^"]*"[^>]*>(.*?)</span>', re.DOTALL
)
# 블록 span 열기
SPAN_OPEN_RE = re.compile(r'<span\s+style="color:\s*([^";]+)[^"]*"[^>]*>\s*$')
# span 닫기
SPAN_CLOSE_RE = re.compile(r'^\s*</span>\s*$')

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
HR_RE = re.compile(r"^---+\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
BLOCKQUOTE_RE = re.compile(r"^>\s*(.*)$")
NAV_LINK_RE = re.compile(r"^\[←.*\]|^\[.*→\]")


def parse_color(color_str):
    """색상 문자열을 RGBColor로 변환."""
    color_str = color_str.strip().lower()
    return COLOR_MAP.get(color_str)


def add_rich_text(paragraph, text, base_font=FONT_BODY, base_size=BODY_SIZE, base_color=None):
    """텍스트에서 인라인 span(색상), **bold**, 링크를 파싱하여 run으로 추가."""
    text = COMMENT_RE.sub("", text)

    # 인라인 span 파싱 — 색상 적용
    pos = 0
    for m in INLINE_SPAN_RE.finditer(text):
        # span 이전 텍스트 (base_color 적용)
        before = text[pos:m.start()]
        if before:
            _add_formatted_runs(paragraph, before, base_font, base_size, color=base_color)
        # span 내부 텍스트 (span 색상 적용 + bold 강제)
        span_color = parse_color(m.group(1))
        inner = m.group(2)
        _add_formatted_runs(paragraph, inner, base_font, base_size, color=span_color, force_bold=True)
        pos = m.end()

    # 남은 텍스트
    remaining = text[pos:]
    if remaining:
        # 혹시 남은 span 태그 잔여물 정리
        remaining = re.sub(r'<span[^>]*>', '', remaining)
        remaining = remaining.replace('</span>', '')
        _add_formatted_runs(paragraph, remaining, base_font, base_size, color=base_color)


def _add_formatted_runs(paragraph, text, font_name, font_size, color=None, force_bold=False):
    """**bold** 마크다운과 [링크](url)를 파싱하여 run으로 추가."""
    # 링크 제거 (텍스트만 유지)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # ⚠️ 태그 제거
    text = re.sub(r'\[⚠️[^\]]*\]', '', text)

    pos = 0
    for m in BOLD_RE.finditer(text):
        before = text[pos : m.start()]
        if before:
            run = paragraph.add_run(before)
            set_font(run, font_name, font_size, bold=force_bold, color=color)
        bold_text = m.group(1)
        run = paragraph.add_run(bold_text)
        set_font(run, font_name, font_size, bold=True, color=color)
        pos = m.end()
    remaining = text[pos:]
    if remaining:
        run = paragraph.add_run(remaining)
        set_font(run, font_name, font_size, bold=force_bold, color=color)


def add_image(doc, img_path, base_dir):
    """이미지를 InlineShape로 삽입. 파일 없으면 플레이스홀더 텍스트."""
    full_path = Path(base_dir) / img_path
    if full_path.exists():
        MAX_W = 6.3   # inches
        MAX_H = 6.0   # inches (이미지 아래 공백 최소화)
        try:
            pil_img = PILImage.open(str(full_path))
            w_px, h_px = pil_img.size
            h_at_max_w = h_px / w_px * MAX_W
            if h_at_max_w > MAX_H:
                doc.add_picture(str(full_path), height=Inches(MAX_H))
            else:
                doc.add_picture(str(full_path), width=Inches(MAX_W))
        except Exception:
            doc.add_picture(str(full_path), width=Inches(MAX_W))
        p = doc.paragraphs[-1]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[이미지 플레이스홀더: {img_path}]")
        set_font(run, color=RGBColor(0x99, 0x99, 0x99))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_cell_border(cell, **kwargs):
    """셀 테두리 설정. kwargs: top, bottom, left, right = (size, color)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tc.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for side, (sz, color) in kwargs.items():
        el = tcBorders.find(qn(f"w:{side}"))
        if el is None:
            el = tc.makeelement(qn(f"w:{side}"), {})
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")


def _set_cell_padding(cell, top=80, bottom=80, left=100, right=100):
    """셀 패딩 설정 (단위: twips)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = tc.makeelement(qn("w:tcMar"), {})
        tcPr.append(tcMar)
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = tc.makeelement(qn(f"w:{side}"), {})
            tcMar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def _set_cell_valign(cell, align="center"):
    """셀 수직 정렬."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    va = tcPr.find(qn("w:vAlign"))
    if va is None:
        va = tc.makeelement(qn("w:vAlign"), {})
        tcPr.append(va)
    va.set(qn("w:val"), align)


def _add_cell_rich_text(paragraph, text, font_name=FONT_TABLE, font_size=TABLE_SIZE):
    """테이블 셀 내부 텍스트에서 **bold**와 <span color:red> 파싱."""
    text = re.sub(r'<span\s+style="color:\s*([^";]+)[^"]*"[^>]*>', r'{{COLOR:\1}}', text)
    text = text.replace('</span>', '{{/COLOR}}')

    parts = re.split(r'(\{\{COLOR:[^}]+\}\}|\{\{/COLOR\}\})', text)
    current_color = None

    for part in parts:
        if part.startswith('{{COLOR:'):
            color_str = part[8:-2].strip().lower()
            current_color = COLOR_MAP.get(color_str)
            continue
        elif part == '{{/COLOR}}':
            current_color = None
            continue

        if not part:
            continue

        pos = 0
        for m in BOLD_RE.finditer(part):
            before = part[pos:m.start()]
            if before:
                run = paragraph.add_run(before)
                set_font(run, font_name, font_size, color=current_color)
            bold_text = m.group(1)
            run = paragraph.add_run(bold_text)
            set_font(run, font_name, font_size, bold=True, color=current_color)
            pos = m.end()
        remaining = part[pos:]
        if remaining:
            run = paragraph.add_run(remaining)
            set_font(run, font_name, font_size, color=current_color)


def add_table(doc, rows):
    """마크다운 파이프 테이블을 고급 스타일 docx Table로 변환."""
    if len(rows) < 2:
        return
    header_parts = rows[0].split("|")
    headers = [c.strip() for c in header_parts[1:-1]]  # 앞뒤 빈 요소 제거, 중간 빈 셀 유지
    data_rows = []
    for r in rows[2:]:
        parts = r.split("|")
        cells = [c.strip() for c in parts[1:-1]]  # 앞뒤 빈 요소 제거, 중간 빈 셀 유지
        data_rows.append(cells)

    ncols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    # ─── 헤더 행 ───
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean_h = re.sub(r"\*\*(.+?)\*\*", r"\1", h)
        clean_h = INLINE_SPAN_RE.sub(r"\2", clean_h)
        run = p.add_run(clean_h)
        set_font(run, FONT_TABLE, Pt(11), bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 헤더 배경: 진한 블루
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): "0062B8", qn("w:val"): "clear"})
        shading.append(shd)
        # 셀 패딩 + 수직 중앙
        _set_cell_padding(cell, 60, 60, 100, 100)
        _set_cell_valign(cell, "center")
        # 테두리: 검정 0.5pt (디딤돌 양식) + 헤더 하단 1pt
        _set_cell_border(cell, top=(4, "000000"), bottom=(8, "000000"), left=(4, "000000"), right=(4, "000000"))

    # ─── 데이터 행 ───
    for ri, row_data in enumerate(data_rows):
        for ci in range(min(len(row_data), ncols)):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            # 리치 텍스트 파싱 (bold, 빨간색 유지)
            _add_cell_rich_text(p, row_data[ci])
            # 셀 패딩 + 수직 중앙
            _set_cell_padding(cell, 60, 60, 100, 100)
            _set_cell_valign(cell, "center")
            # 교차 행 배경
            bg_color = "F4F7FC" if ri % 2 == 0 else "FFFFFF"
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): bg_color, qn("w:val"): "clear"})
            shading.append(shd)
            # 테두리: 검정 0.5pt (디딤돌 양식)
            _set_cell_border(cell, top=(4, "000000"), bottom=(4, "000000"), left=(4, "000000"), right=(4, "000000"))


def convert_md_to_docx(md_path, docx_path):
    """단일 md 파일을 docx로 변환."""
    md_path = Path(md_path)
    base_dir = md_path.parent

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_style_defaults(doc)
    # 디딤돌 양식 여백: 상하 2.54cm, 좌우 2.54cm + 페이지 번호
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        # 페이지 번호 (하단 중앙)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = run2._element.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._element.append(fldChar2)
        set_font(run, FONT_BODY, Pt(9))

    i = 0
    table_buffer = []
    in_table = False
    current_color = None  # 멀티라인 span 색상 추적
    consecutive_blank = 0  # 연속 빈 줄 카운트
    has_content = False  # 본문 내용이 시작되었는지

    # 마지막 본문 줄 인덱스 계산 (끝 부분 --- 스킵용)
    last_content_idx = len(lines) - 1
    while last_content_idx >= 0:
        l = lines[last_content_idx].strip()
        if l and not l.startswith("<!--") and not l.startswith("*") and l != "---" and not l.startswith("[←"):
            break
        last_content_idx -= 1

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # 주석 전용 줄 스킵
        if line.strip().startswith("<!--") and line.strip().endswith("-->"):
            i += 1
            continue

        # 네비게이션 링크 스킵 ([← 인덱스] 등)
        if NAV_LINK_RE.match(line.strip()):
            i += 1
            continue

        # 블록 span 열기 감지
        span_open = SPAN_OPEN_RE.match(line.strip())
        if span_open:
            current_color = parse_color(span_open.group(1))
            i += 1
            continue

        # span 닫기 감지
        if SPAN_CLOSE_RE.match(line.strip()):
            current_color = None
            i += 1
            continue

        # CEO 입력 필요 / ⚠️ 태그만 있는 줄 스킵
        if line.strip().startswith("[CEO 입력 필요") or line.strip().startswith("[⚠️"):
            i += 1
            continue

        # 빈 줄 처리 — 연속 빈 줄 최대 1개만 허용
        if not line.strip():
            if in_table and table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False
                doc.add_paragraph()  # 테이블 뒤 구분 단락
            consecutive_blank += 1
            i += 1
            continue
        else:
            consecutive_blank = 0

        # 테이블 행
        if TABLE_ROW_RE.match(line):
            if TABLE_SEP_RE.match(line):
                table_buffer.append(line)
            else:
                table_buffer.append(line)
            in_table = True
            i += 1
            continue
        elif in_table and table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        # 수평선 처리 — 첫/마지막 근처는 스킵, 본문 중간은 얇은 구분선
        if HR_RE.match(line):
            if not has_content or i > last_content_idx:
                # 본문 시작 전 또는 끝 부분 --- → 스킵 (빈 페이지 방지)
                i += 1
                continue
            # 본문 중간 --- → 얇은 수평선 (페이지 브레이크 아님)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, Pt(4), Pt(4))
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 문서 끝 메타 (*v1.3...) 스킵
        if line.strip().startswith("*") and line.strip().endswith("*") and len(line.strip()) > 5:
            i += 1
            continue

        # 제목
        hm = HEADING_RE.match(line)
        if hm:
            has_content = True
            level = len(hm.group(1))
            title_text = hm.group(2).strip()
            title_text = re.sub(r"\*\*(.+?)\*\*", r"\1", title_text)
            title_text = INLINE_SPAN_RE.sub(r"\2", title_text)
            heading_level = min(level, 3)
            h = doc.add_heading(title_text, level=heading_level)
            for run in h.runs:
                set_font(
                    run,
                    FONT_TITLE,
                    [TITLE_SIZE, H1_SIZE, H2_SIZE, H3_SIZE][heading_level],
                    bold=True,
                )
            i += 1
            continue

        # 이미지
        img_m = IMG_RE.search(line)
        if img_m and line.strip().startswith("!"):
            add_image(doc, img_m.group(2), base_dir)
            if img_m.group(1):
                cap = doc.add_paragraph()
                run = cap.add_run(img_m.group(1))
                set_font(run, size=CAPTION_SIZE, color=RGBColor(0x8E, 0x8E, 0x8E))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(cap, Pt(0), Pt(2))
            i += 1
            continue

        # 인라인 span (한 줄에 열고 닫는 경우 — 주로 파란색 작성요령)
        inline_span = INLINE_SPAN_RE.match(line.strip())
        if inline_span and '<span' in line and '</span>' in line:
            color = parse_color(inline_span.group(1))
            inner_text = inline_span.group(2)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, Pt(4), Pt(4))
            add_rich_text(p, inner_text, base_color=color)
            i += 1
            continue

        # 블록인용
        bq_m = BLOCKQUOTE_RE.match(line)
        if bq_m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            set_paragraph_spacing(p, Pt(2), Pt(2))
            add_rich_text(p, bq_m.group(1), base_size=Pt(10), base_color=current_color)
            i += 1
            continue

        # 일반 텍스트 — current_color 적용 + 들여쓰기 레벨 감지
        p = doc.add_paragraph()
        set_paragraph_spacing(p)
        stripped = line.strip()

        # 마크다운 들여쓰기 감지 (4칸 = 1레벨)
        leading_spaces = len(line) - len(line.lstrip())
        indent_level = leading_spaces // 4

        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(0.8 * indent_level)
        else:
            # 불릿·번호 항목이 아닌 일반 서술문만 첫줄 들여쓰기 (디딤돌 양식)
            is_bullet = stripped and stripped[0] in '‣→⇒①②③④⑤⑥⑦⑧⑨-'
            is_numbered = bool(re.match(r'^\d+[\.\)]\s', stripped))
            if not is_bullet and not is_numbered and len(stripped) > 20:
                p.paragraph_format.first_line_indent = Cm(0.5)
        add_rich_text(p, line, base_color=current_color)
        i += 1

    # 남은 테이블 버퍼
    if table_buffer:
        add_table(doc, table_buffer)

    doc.save(str(docx_path))
    print(f"✅ {md_path.name} → {Path(docx_path).name}")


def main():
    if len(sys.argv) < 2:
        print("사용법: python3 md-to-docx.py input.md [output.docx]")
        print("        python3 md-to-docx.py --dir /path/to/folder/")
        sys.exit(1)

    if sys.argv[1] == "--dir":
        folder = Path(sys.argv[2])
        for md_file in sorted(folder.glob("chapter*.md")):
            docx_file = md_file.with_suffix(".docx")
            convert_md_to_docx(md_file, docx_file)
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace(".md", ".docx")
        convert_md_to_docx(md_file, docx_file)


if __name__ == "__main__":
    main()
